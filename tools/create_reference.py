
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

path ="reference.docx"

doc = Document()


styles = doc.styles

for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
    style = styles[style_name]
    style.font.name = "Vazirmatn"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Vazirmatn")

styles["Normal"].font.size = Pt(11)
styles["Title"].font.size = Pt(22)
styles["Heading 1"].font.size = Pt(16)
styles["Heading 2"].font.size = Pt(14)
styles["Heading 3"].font.size = Pt(12)


title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Frontend Engineering Handbook")

doc.add_paragraph(
    "این فایل به عنوان reference.docx برای تبدیل Markdown به Word با Pandoc استفاده می‌شود."
)

doc.add_heading("Heading 1 Example", level=1)
doc.add_paragraph("متن نمونه برای بررسی استایل فصل‌ها.")

doc.add_heading("Heading 2 Example", level=2)
doc.add_paragraph("متن نمونه زیرعنوان.")

doc.add_heading("Heading 3 Example", level=3)
doc.add_paragraph("متن نمونه زیر بخش.")

doc.save(path)

path
